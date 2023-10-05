import asyncio
import csv
import os
from copy import deepcopy
from docx import Document
import chardet

def detect_and_read_csv(filename):
    with open(filename, 'rb') as file:
        # 使用chardet检测文件编码
        result = chardet.detect(file.read())

    # 获取检测到的编码
    encoding = result['encoding']

    # 打开文件并根据检测到的编码解码
    with open(filename, 'r', encoding=encoding) as csvfile:
        reader = csv.reader(csvfile)
        data = list(reader)

    return data

async def process_row(row, indices, order, separator):
    # 获取CSV文件中的指定参数作为文件名
    file_name_parts = [row[i] for i in indices]
    ordered_file_name_parts = [file_name_parts[i] for i in order]
    file_name = separator.join(ordered_file_name_parts)

    # 创建新的Word文档
    doc = Document()

    # 复制模板文档中的段落和表格到新文档
    for element in template_doc.element.body:
        new_element = deepcopy(element)
        doc.element.body.append(new_element)

    # 替换文本
    for j, value in enumerate(row):
        if value:
            text = str(value)
        else:
            text = ""
        # 替换文本
        for paragraph in doc.paragraphs:
            paragraph.text = paragraph.text.replace(f"{{{{{headers[j]}}}}}", text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(f"{{{{{headers[j]}}}}}", text)

    # 保存新的Word文档到指定文件夹
    doc.save(os.path.join(output_folder, f'{file_name}.docx'))

if __name__ == '__main__':
    params = {
        'indices': [0, 1,2,3],
        'order': [0, 1,2,3],
        'separator': '_',
        'output_folder': 'output_folder666',
        'csv_file_name': 'output.csv',
        'template_file_name': 'your_template_document.docx',
    }

    if max(params['order']) >= len(params['indices']):
        raise ValueError("Order index is out of range of the indices list.")

    # 使用参数
    output_folder = params['output_folder']
    os.makedirs(output_folder, exist_ok=True)

    # 使用detect_and_read_csv函数来读取CSV文件
    data = detect_and_read_csv(params['csv_file_name'])
    template_doc = Document(params['template_file_name'])
    headers = data[0]

    async def main():
        await asyncio.gather(
            *(process_row(row, params['indices'], params['order'], params['separator']) for row in data[1:]))
            
    asyncio.run(main())
