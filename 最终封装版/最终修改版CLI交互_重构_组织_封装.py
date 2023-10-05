import asyncio
import csv
import os
from copy import deepcopy
from docx import Document
import chardet
import json
import sys

def detect_and_read_csv(filename):
    with open(filename, 'rb') as file:
        result = chardet.detect(file.read())

    encoding = result['encoding']

    with open(filename, 'r', encoding=encoding) as csvfile:
        reader = csv.reader(csvfile)
        data = list(reader)

    return data

async def process_row(row, indices, order, separator, output_folder, template_doc, headers):
    file_name_parts = [row[i] for i in indices]
    ordered_file_name_parts = [file_name_parts[i] for i in order]
    file_name = separator.join(ordered_file_name_parts)

    doc = Document()
    for element in template_doc.element.body:
        new_element = deepcopy(element)
        doc.element.body.append(new_element)

    for j, value in enumerate(row):
        if value:
            text = str(value)
        else:
            text = ""
        for paragraph in doc.paragraphs:
            paragraph.text = paragraph.text.replace(f"{{{{{headers[j]}}}}}", text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(f"{{{{{headers[j]}}}}}", text)

    doc.save(os.path.join(output_folder, f'{file_name}.docx'))

def read_config():
    script_dir = os.path.dirname(os.path.realpath(sys.argv[0]))
    config_file = os.path.join(script_dir, 'config.json')
    
    with open(config_file, 'r') as file:
        config = json.load(file)
        indices = config['indices']
        order = config['order']
        separator = config['separator']
        output_folder = config['output_folder']
        csv_file_name = config['csv_file_name']
        template_file_name = config['template_file_name']

    return indices, order, separator, output_folder, csv_file_name, template_file_name

def main():
    indices, order, separator, output_folder, csv_file_name, template_file_name = read_config()
    
    output_folder = os.path.abspath(output_folder)
    os.makedirs(output_folder, exist_ok=True)
    
    data = detect_and_read_csv(csv_file_name)
    template_doc = Document(template_file_name)
    headers = data[0]

    async def process_data():
        await asyncio.gather(
            *(process_row(row, indices, order, separator, output_folder, template_doc, headers) for row in data[1:])
        )

    asyncio.run(process_data())

if __name__ == '__main__':
    main()
