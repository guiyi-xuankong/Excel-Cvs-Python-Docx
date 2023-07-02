from docx import Document
import csv
import os
from copy import deepcopy

# 创建保存Word文档的文件夹
output_folder = "Excel-Cvs-Python-Docx"
os.makedirs(output_folder, exist_ok=True)

# 读取CSV文件
with open('output.csv', 'r') as csvfile:
    reader = csv.reader(csvfile)
    data = list(reader)
    headers = data[0]  # 获取标题行字段名称

# 读取Word模板文档
template_doc = Document('your_template_document.docx')

# 遍历每一行数据（从第二行开始）
for row in data[1:]:
    # 创建新的Word文档
    doc = Document()

    # 复制模板文档中的段落和表格到新文档
    for element in template_doc.element.body:
        new_element = deepcopy(element)
        doc.element.body.append(new_element)

    # 直接使用参数的值进行替换
    for j, value in enumerate(row):
        if value:
            text = str(value)
        else:
            text = ""
        # 替换文本
        for paragraph in doc.paragraphs:
            paragraph.text = paragraph.text.replace(f"{{{{{headers[j]}}}}}", text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(f"{{{{{headers[j]}}}}}", text)

    # 获取CSV文件中的第一个参数作为文件名
    file_name = row[0]

    # 保存新的Word文档到指定文件夹
    doc.save(os.path.join(output_folder, f'{file_name}.docx'))

# 不再需要关闭模板文档